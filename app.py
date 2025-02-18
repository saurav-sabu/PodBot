import streamlit as st
import pyttsx3
from src.helper import *
from docx import Document  # Import for generating .docx files
from io import BytesIO
from gtts import gTTS
import tempfile 
import urllib.request  # Ensure urllib.request is available
import re
__import__('pysqlite3')
import sys
sys.modules['sqlite3'] = sys.modules.pop('pysqlite3')

st.set_page_config(
    page_title="Podcast Generator",  # App Name
    page_icon="🎙️",  # Icon (You can use an emoji or a favicon URL)
    layout="wide"  # Optional: Makes the layout wider
)

engine = pyttsx3.init()

def clean_text_for_speech(text):
    """
    Removes formatting symbols like **bold**, __underline__, and *italic* 
    before passing the text to pyttsx3.
    """
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)  # Remove **bold**
    text = re.sub(r"__(.*?)__", r"\1", text)      # Remove __underline__
    text = re.sub(r"\*(.*?)\*", r"\1", text)      # Remove *italic*
    return text

# Function for offline speech synthesis (fixes disappearing issue)
def generate_audio(text):
    clean_text = clean_text_for_speech(text)  # Clean the text before speaking
    
    buffer = BytesIO()  
    engine.save_to_file(clean_text, "temp_audio.mp3")  
    engine.runAndWait()

    with open("temp_audio.mp3", "rb") as f:
        buffer.write(f.read())

    buffer.seek(0)  
    return buffer


# Function to generate a .docx file
def generate_docx(text):
    doc = Document()
    doc.add_heading("Podcast Transcript", level=1)
    doc.add_paragraph(text)
    
    # Save the document in-memory
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Custom Styling
st.markdown(
    """
    <style>
    .main { background-color: #f4f4f4; }
    .stTextInput, .stButton, .stTextArea { font-size: 18px !important; }
    h1, h2, h3, h4 { text-align: center; color: #ff5733; }
    .sidebar .sidebar-content { background-color: #2c3e50; color: white; }
    </style>
    """,
    unsafe_allow_html=True
)


# Page Navigation
st.sidebar.title("🎙️ Podcast App")
st.sidebar.markdown("---")
page = st.sidebar.radio("📌 Navigate to", ["🏠 Home", "ℹ️ About", "📧 Contact"])

if page == "🏠 Home":
    st.title("🎤 Podcast Generator")
    st.markdown("### 🎯 Enter a topic to generate an engaging podcast")
    topic = st.text_input("🔍 Topic:", "Machine Learning")
    llm = initialize_llm()
    if st.button("🚀 Generate Podcast"):
        with st.spinner("⏳ Generating podcast..."):
            interviewer, guest = Agents(topic, llm)
            interview_task, guest_task = Tasks(interviewer, guest,topic)
            st.subheader("📝 Podcast Transcript:")
            with st.expander("📜 View Transcript"):
                text = podcast_generation(interviewer, guest, interview_task, guest_task,topic)
                st.write(text)
                # text_to_speech(text)  

            docx_file = generate_docx(text)
            st.download_button(
                label="📥 Download Transcript",
                data=docx_file,
                file_name=f"Podcast_Transcript_{topic.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            # Generate and Play Audio using gTTS
            st.subheader("🔊 Listen to Podcast")
            st.markdown("Click below to hear the podcast transcript read aloud.")
            audio_buffer = generate_audio(text)
            st.audio(audio_buffer, format="audio/mp3")

elif page == "ℹ️ About":
    st.title("📖 About the App")
    st.markdown(
        """
        🎙️ **Podcast Generator** is an AI-powered application that creates engaging podcast-like content based on your chosen topic.
        
        🔹 **Features:**
        - 📢 Generate structured podcast transcripts
        - 🎧 Text-to-Speech functionality
        - 📜 Beautiful UI with interactive elements
        
        🚀 **Enjoy creating podcasts effortlessly!**
        """
    )

elif page == "📧 Contact":
    st.title("📞 Contact Us")
    st.markdown("✉️ Reach out to us at: [contact@podcastapp.com](mailto:contact@podcastapp.com)")
    st.markdown("📌 Follow us on [Twitter](https://twitter.com/podcastapp) for updates!")
